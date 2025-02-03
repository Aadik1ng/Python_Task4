import os
import io
import pandas as pd
import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

# Constants
OUTPUT_DIR = "output"
LOGO_FILENAME = "logo.png"  # Place your company logo in the same folder as this script

# Ensure the output folder exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_word_doc(row):
    """
    Generate a personalized Word document for the given row data.
    """
    try:
        doc = Document()

        # Add company logo if available
        if os.path.exists(LOGO_FILENAME):
            try:
                doc.add_picture(LOGO_FILENAME, width=Inches(2))
            except Exception as e:
                st.warning(f"Could not add logo to Word document: {e}")

        # Add heading and content with placeholders replaced by provided data
        doc.add_heading("Welcome to Your New Role!", level=1)
        name = row.get("Name", "Employee")
        position = row.get("Position", "your position")
        company = row.get("Company Name", "our company")
        joining_date = row.get("Joining Date", "TBA")
        
        doc.add_paragraph(f"Dear {name},")
        doc.add_paragraph(
            f"We are delighted to welcome you as our new {position} at {company}. "
            f"Your joining date is {joining_date}."
        )
        doc.add_paragraph("We are confident that you will make a significant contribution to our team.")
        doc.add_paragraph("Best regards,")
        doc.add_paragraph("Human Resources")
        
        # Create a safe filename (e.g. John_Doe_CompanyX.docx)
        safe_name = name.replace(" ", "_")
        safe_company = company.replace(" ", "_")
        filename = os.path.join(OUTPUT_DIR, f"{safe_name}_{safe_company}.docx")
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error generating Word document for {row.get('Name', 'Employee')}: {e}")
        return None

from PIL import Image

def generate_pdf_doc(row):
    try:
        name = row.get("Name", "Employee")
        position = row.get("Position", "your position")
        company = row.get("Company Name", "our company")
        joining_date = row.get("Joining Date", "TBA")

        safe_name = name.replace(" ", "_")
        safe_company = company.replace(" ", "_")
        filename = os.path.join(OUTPUT_DIR, f"{safe_name}_{safe_company}.pdf")

        c = canvas.Canvas(filename, pagesize=LETTER)
        width, height = LETTER

        y_position = height - 50

        # Draw logo at the top with correct aspect ratio
        if os.path.exists(LOGO_FILENAME):
            try:
                img = Image.open(LOGO_FILENAME)
                img_width, img_height = img.size

                # Define max width and height for the logo
                max_width = 150  # Adjust as needed
                max_height = 100  # Adjust as needed

                # Maintain aspect ratio
                ratio = min(max_width / img_width, max_height / img_height)
                new_width = int(img_width * ratio)
                new_height = int(img_height * ratio)

                # Center logo
                x_position = (width - new_width) // 2
                c.drawImage(LOGO_FILENAME, x_position, y_position - new_height, width=new_width, height=new_height)
                y_position -= new_height + 20  # Add some space below the logo
            except Exception as e:
                st.warning(f"Could not add logo to PDF: {e}")

        c.setFont("Helvetica-Bold", 20)
        c.drawString(50, y_position, "Welcome to Your New Role!")
        y_position -= 40

        c.setFont("Helvetica", 12)
        c.drawString(50, y_position, f"Dear {name},")
        y_position -= 25

        # Handle multi-line text
        message = (
            f"We are delighted to welcome you as our new {position} at {company}. "
            f"Your joining date is {joining_date}. "
            "We are confident that you will make a significant contribution to our team."
        )
        text_object = c.beginText(50, y_position)
        text_object.setFont("Helvetica", 12)
        for line in message.split(". "):
            if not line.endswith("."):
                line += "."
            text_object.textLine(line)
        c.drawText(text_object)
        y_position -= 80

        c.drawString(50, y_position, "Best regards,")
        y_position -= 20
        c.drawString(50, y_position, "Human Resources")

        c.showPage()
        c.save()
        return filename
    except Exception as e:
        st.error(f"Error generating PDF document for {row.get('Name', 'Employee')}: {e}")
        return None



def process_data(data_rows, export_pdf, export_word):
    """
    Process the provided data rows (each a dictionary) and generate documents for each.
    """
    generated_files = []
    for row in data_rows:
        # Optionally convert joining date to a string (if it's a datetime)
        joining_date = row.get("Joining Date")
        if isinstance(joining_date, (pd.Timestamp, datetime)):
            row["Joining Date"] = joining_date.strftime("%Y-%m-%d")
        
        # Generate Word document if selected
        if export_word:
            word_file = generate_word_doc(row)
            if word_file:
                generated_files.append(word_file)
        # Generate PDF document if selected
        if export_pdf:
            pdf_file = generate_pdf_doc(row)
            if pdf_file:
                generated_files.append(pdf_file)
    return generated_files

def process_excel(file_bytes, export_pdf, export_word):
    """
    Process the uploaded Excel file and generate documents for each row.
    """
    try:
        # Read the Excel file from bytes
        df = pd.read_excel(io.BytesIO(file_bytes))
    except Exception as e:
        st.error(f"Could not read Excel file: {e}")
        return []

    # Required columns
    required_columns = {"Name", "Email", "Company Name", "Position", "Joining Date"}
    if not required_columns.issubset(set(df.columns)):
        st.error(f"Excel file must contain the columns: {', '.join(required_columns)}")
        return []

    data_rows = []
    for index, row in df.iterrows():
        row_data = row.to_dict()
        data_rows.append(row_data)
    return process_data(data_rows, export_pdf, export_word)

def manual_entry_form():
    """
    Display a form for manual entry and return a dictionary of the entered data.
    """
    with st.form(key="manual_entry_form", clear_on_submit=True):
        name = st.text_input("Name")
        email = st.text_input("Email")
        company = st.text_input("Company Name")
        position = st.text_input("Position")
        joining_date = st.date_input("Joining Date")
        submitted = st.form_submit_button("Add Entry")
        if submitted:
            # Convert the date to string format if necessary
            entry = {
                "Name": name or "Employee",
                "Email": email or "",
                "Company Name": company or "our company",
                "Position": position or "your position",
                "Joining Date": joining_date.strftime("%Y-%m-%d")
            }
            return entry
    return None

def main():
    st.title("Personalized Document Generator")
    st.write("Generate personalized Word or PDF documents based on employee data.")

    # Option selectors for output formats
    col1, col2 = st.columns(2)
    with col1:
        export_pdf = st.checkbox("Export to PDF", value=True)
    with col2:
        export_word = st.checkbox("Export to Word", value=True)

    # Allow user to choose the input method
    input_method = st.radio("Select input method:", ("Upload Excel File", "Manual Entry"))

    generated_files = []
    data_rows = []

    if input_method == "Upload Excel File":
        st.write("Upload one or more Excel files (xlsx format) with the required columns:")
        # Allow multiple file uploads
        uploaded_files = st.file_uploader("Choose Excel file(s)", type=["xlsx", "xls"], accept_multiple_files=True)
        if uploaded_files:
            if st.button("Generate Documents", key="excel_generate"):
                # Process each file
                for uploaded_file in uploaded_files:
                    file_bytes = uploaded_file.read()
                    files_from_excel = process_excel(file_bytes, export_pdf, export_word)
                    generated_files.extend(files_from_excel)
                if generated_files:
                    st.success("Documents generated successfully!")
                    st.write("Files saved in the `output/` folder:")
                    for file in generated_files:
                        st.write(f"- {file}")
                else:
                    st.error("No documents were generated.")
    else:
        st.subheader("Manual Entry of Employee Data")
        # Initialize session state for storing manual entries
        if "manual_entries" not in st.session_state:
            st.session_state.manual_entries = []
        
        new_entry = manual_entry_form()
        if new_entry:
            st.session_state.manual_entries.append(new_entry)
            st.success("Entry added successfully!")

        if st.session_state.manual_entries:
            st.write("### Current Entries")
            df_entries = pd.DataFrame(st.session_state.manual_entries)
            st.dataframe(df_entries)
            
            # Option to clear all entries
            if st.button("Clear All Entries"):
                st.session_state.manual_entries = []
                st.experimental_rerun()
            
            if st.button("Generate Documents", key="manual_generate"):
                data_rows = st.session_state.manual_entries
                generated_files = process_data(data_rows, export_pdf, export_word)
                if generated_files:
                    st.success("Documents generated successfully!")
                    st.write("Files saved in the `output/` folder:")
                    for file in generated_files:
                        st.write(f"- {file}")
                else:
                    st.error("No documents were generated.")

if __name__ == "__main__":
    main()
